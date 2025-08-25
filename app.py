import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from docx.oxml import OxmlElement
import anthropic
import json
from docx.shared import RGBColor
import unicodedata
from datetime import datetime
import re

def normaliser_texte(text):
    # Normalisation Unicode (accents)
    texte = unicodedata.normalize('NFC', text)
    
    # Remplacement des apostrophes typographiques par apostrophe simple
    texte = texte.replace("’", "'").replace("‘", "'").replace("‛", "'")

    # Remplacer les espaces spéciaux par un espace classique
    espaces_speciaux = ['\xa0', '\u202f', '\u2009', '\u200A', '\t', '\r', '\n']
    for esp in espaces_speciaux:
        texte = texte.replace(esp, ' ')

    # Supprimer espaces début/fin
    texte = texte.strip()

    return texte

def get_runs_from_element(elem):
    return [r for r in elem if r.tag == qn('w:r')]

def runs_equivalents(r1, r2):
    def get_bool(el, tag):
        val = el.find(qn(f"w:{tag}"))
        return val is not None and val.get(qn('w:val'), 'true') in ('true', '1')

    def get_size(el):
        sz = el.find(qn("w:rPr"))
        if sz is not None:
            size = sz.find(qn("w:sz"))
            if size is not None:
                return size.get(qn("w:val"))
        return None

    r1rpr = r1.find(qn("w:rPr"))
    r2rpr = r2.find(qn("w:rPr"))
    if r1rpr is None or r2rpr is None:
        return False

    return (
        get_bool(r1rpr, "b") == get_bool(r2rpr, "b") and
        get_bool(r1rpr, "i") == get_bool(r2rpr, "i") and
        get_bool(r1rpr, "u") == get_bool(r2rpr, "u") and
        get_size(r1) == get_size(r2)
    )

def fusionner_runs_xml(runs):
    if not runs:
        return []

    fusionnes = []
    current_run = runs[0]
    current_text = current_run.find(qn("w:t")).text if current_run.find(qn("w:t")) is not None else ""

    for run in runs[1:]:
        r_text_el = run.find(qn("w:t"))
        r_text = r_text_el.text if r_text_el is not None else ""

        if runs_equivalents(current_run, run):
            current_text += r_text
        else:
            # Finaliser le run courant
            t_el = current_run.find(qn("w:t"))
            if t_el is not None:
                t_el.text = current_text
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            fusionnes.append(current_run)
            current_run = run
            current_text = r_text

    # Dernier run
    t_el = current_run.find(qn("w:t"))
    if t_el is not None:
        t_el.text = current_text
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    fusionnes.append(current_run)

    return fusionnes

def fusionner_runs_paragraphe(p: Paragraph):
    p_xml = p._element
    nouveaux_enfants = []

    enfants = list(p_xml)  # snapshot pour éviter modification en cours
    buffer_runs = []

    def vider_buffer():
        nonlocal buffer_runs
        if buffer_runs:
            fusionnes = fusionner_runs_xml(buffer_runs)
            nouveaux_enfants.extend(fusionnes)
            buffer_runs = []

    for el in enfants:
        if el.tag == qn("w:r"):
            buffer_runs.append(el)
        elif el.tag == qn("w:hyperlink"):
            vider_buffer()
            nouveaux_enfants.append(el)
        else:
            vider_buffer()
            nouveaux_enfants.append(el)

    vider_buffer()  # en fin de paragraphe

    # Remplacer tous les enfants par ceux reconstruits
    p_xml.clear()
    for child in nouveaux_enfants:
        p_xml.append(child)

def fusionner_runs_cellule(cell: _Cell):
    for paragraph in cell.paragraphs:
        fusionner_runs_paragraphe(paragraph)

def fusionner_runs_similaires(doc: Document):
    for paragraph in doc.paragraphs:
        fusionner_runs_paragraphe(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fusionner_runs_cellule(cell)

def get_text(doc) :
	text = ""
	for para in doc.paragraphs :
		text += para.text
	for table in doc.tables :
		for row in table.rows :
			for cell in row.cells :
				text += cell.text
	return text

def get_claude_response(text) :
	api_key = st.secrets["ANTHROPIC_API_KEY"]
	client = anthropic.Anthropic(api_key=api_key)


	message_content = f"""
	You are tasked with correcting spelling mistakes, grammatical errors, and anglicisms in a French text. Your goal is to provide a dictionary of corrections made, without returning the corrected text itself.

	Here is the text to be corrected:
	<text_to_correct>
	{text}
	</text_to_correct>

	Follow these steps to complete the task:

	1. Carefully read through the text, identifying any spelling mistakes, grammatical errors, and anglicisms.

	2. For each error you identify:
	   a. Note the incorrect word or phrase
	   b. Determine the correct French equivalent

	3. Create a dictionary where:
	   - The keys are the incorrect words or phrases
	   - The values are the correct French equivalents

	4. When correcting:
	   - Fix spelling errors according to standard French orthography
	   - Correct grammatical mistakes to ensure proper French syntax
	   - Replace anglicisms with appropriate French terms or expressions

	5. If a word or phrase could have multiple correct alternatives, choose the most appropriate one based on the context.

	6. Do not include any explanations or comments in the dictionary. It should contain only the corrections.

	7. Do not correct proper nouns.

	8. Include several words in the error detected to distinguish it from other parts of the text.

	9. If no errors are found, return an empty dictionary.

	Provide your answer in the following format:
	{{
	    "incorrect1": "correct1",
	    "incorrect2": "correct2",
	    ...
	}}

	Remember to return only the JSON dictionary of corrections, without the corrected text or any additional explanations.
	"""

	message = client.messages.create(
		model="claude-sonnet-4-20250514",
		max_tokens=1024,
		messages=[
			{
			"role": "user",
			"content": message_content
			}
		],
		temperature=0
	)
	result = message.content[0].text
	d = json.loads(result)
	return {error: correction for error, correction in d.items() if error != correction}


def split_run_on_error(run, error_text, corr):
    """
    Scinde un run contenant une erreur en trois runs : avant, erreur, après.
    Garde le style original, colore l'erreur en rouge.
    """
    full_text = run.text
    idx = full_text.find(error_text)

    if idx == -1:
        return  # Pas d'erreur trouvée dans ce run

    before = full_text[:idx]
    error = full_text[idx:idx + len(error_text)]
    after = " " + corr + full_text[idx + len(error_text):]

    paragraph = run._parent  # ✅ Correct : le paragraphe contenant ce run
    r_index = list(paragraph._element).index(run._element)

    # Supprimer l'ancien run
    paragraph._element.remove(run._element)

    # Crée les trois nouveaux runs
    def clone_style(source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    new_runs = []

    if before:
        run_before = paragraph.add_run(before)
        clone_style(run, run_before)
        new_runs.append(run_before)

    run_error = paragraph.add_run(error)
    clone_style(run, run_error)
    run_error.font.color.rgb = RGBColor(255, 0, 0)  # rouge
    new_runs.append(run_error)

    if after:
        run_after = paragraph.add_run(after)
        clone_style(run, run_after)
        new_runs.append(run_after)

    # Réinsérer les nouveaux runs à la bonne position
    for i, new_run in enumerate(new_runs):
        paragraph._element.insert(r_index + i, new_run._element)

def corriger_fichier(doc, dico) :
	for err in dico :
		# Traitement des paragraphes
		for para in doc.paragraphs :
			for run in para.runs :
				if err in normaliser_texte(run.text) :
					split_run_on_error(run, err, dico[err])

		# Traitement des tableaux
		for table in doc.tables :
			for row in table.rows :
				for cell in row.cells :
					for para in cell.paragraphs :
						for run in para.runs :
							if err in normaliser_texte(run.text) :
								print(err)
								split_run_on_error(run, err, dico[err])

def ouvrir_fichier():
    chemin = filedialog.askopenfilename(
        title="Choisir un fichier Word",
        filetypes=[("Fichiers Word", "*.docx")]
    )
    if chemin:
        try:
            doc = Document(chemin)
            fusionner_runs_similaires(doc)
            text_full = get_text(doc)
            dico = get_claude_response(text_full)
            corriger_fichier(doc, dico)
            doc.save(chemin)
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible de lire le fichier.\n{e}")

# Fenêtre principale
st.title("📄 Correcteur de fichiers Word avec Claude")
uploaded_file = st.file_uploader("Choisissez un fichier Word (.docx)", type=["docx"])

if uploaded_file:
    if st.button("Corriger le fichier"):
        try:
            doc = Document(uploaded_file)
            text_full = get_text(doc)
            dico = get_claude_response(text_full)
            corriger_fichier(doc, dico)

            base_name = uploaded_file.name.replace(".docx", "")
            output_name = f"{base_name}_corrige.docx"

            # Sauvegarde temporaire
            doc.save(output_name)

            with open(output_name, "rb") as f:
                st.download_button(
                    "⬇️ Télécharger le fichier corrigé",
                    f,
                    file_name=output_name
                )

            st.success(f"✅ Fichier corrigé avec succès : {output_name}")

        except Exception as e:
            st.error(f"Erreur : {e}")




